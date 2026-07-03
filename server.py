import os
import certifi

# 彻底清除系统中的 Conda 环境变量干扰，确保只依赖独立的 .venv 虚拟环境
for key in list(os.environ.keys()):
    if "CONDA" in key.upper():
        del os.environ[key]

# 针对 Windows 中文路径 (如 "项目") 导致的 OpenSSL 证书加载 FileNotFoundError 异常进行修复
for env_var in ["SSL_CERT_FILE", "REQUESTS_CA_BUNDLE", "CURL_CA_BUNDLE"]:
    if env_var in os.environ:
        val = os.environ[env_var]
        # 如果路径不存在，或者包含非 ASCII 字符（如中文），则重置为 certifi 的安全路径
        if not os.path.exists(val) or any(ord(c) > 127 for c in val):
            os.environ[env_var] = certifi.where()
    else:
        # 默认注入安全证书，防范潜在的 SSL 报错
        os.environ[env_var] = certifi.where()

# 清理可能导致乱码的证书目录变量
if "SSL_CERT_DIR" in os.environ:
    val = os.environ["SSL_CERT_DIR"]
    if not os.path.exists(val) or any(ord(c) > 127 for c in val):
        del os.environ["SSL_CERT_DIR"]

import io
import re
import json
import uuid
import shutil
import datetime
import logging
import threading
import asyncio
import concurrent.futures
import pandas as pd
import zipfile
from typing import Dict, List, Any, Optional
from contextlib import asynccontextmanager
from fastapi import FastAPI, BackgroundTasks, HTTPException, Body, UploadFile, File, Form
from fastapi.responses import HTMLResponse, StreamingResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from sse_starlette.sse import EventSourceResponse
import webbrowser

from scraper import WebScraper
from slicer import smart_slice
from analyzer import ImageAnalyzer

# 日志配置
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%H:%M:%S"
)
logger = logging.getLogger("Web_OCR_Server")

def open_browser():
    """等待服务器启动后，自动在默认浏览器中打开管理页面"""
    import time
    time.sleep(1.5)
    try:
        webbrowser.open("http://127.0.0.1:8000")
        logger.info("Successfully launched browser to http://127.0.0.1:8000")
    except Exception as e:
        logger.warning(f"Failed to open browser automatically: {e}")

@asynccontextmanager
async def lifespan(app: FastAPI):
    print("\n" + "="*60)
    print("  OmniExtract (万象多模态提取引擎) 服务已成功启动！")
    print("  本地管理页面: http://127.0.0.1:8000")
    print("  请使用浏览器打开上述链接开始进行高精度 OCR/网页提取。")
    print("="*60 + "\n")
    logger.info("OmniExtract 服务初始化就绪。")
    
    # 启动后台线程自动打开浏览器主界面
    threading.Thread(target=open_browser, daemon=True).start()
    yield

app = FastAPI(title="Web OCR Engine API", version="2.0", lifespan=lifespan)

# 启用 CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── 路径常量 ─────────────────────────────────────────────────────────────────
import sys
if getattr(sys, 'frozen', False):
    BASE_DIR = sys._MEIPASS
    DATA_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    DATA_DIR = BASE_DIR

CONFIG_FILE = os.path.join(DATA_DIR, "config.json")
HISTORY_FILE = os.path.join(DATA_DIR, "history.json")
OUTPUT_DIR = os.path.join(DATA_DIR, "output")

# 全局任务状态存储
tasks_db: Dict[str, Dict[str, Any]] = {}
tasks_events: Dict[str, asyncio.Event] = {}

# ── 辅助函数 ─────────────────────────────────────────────────────────────────

def load_configs() -> Dict[str, Dict[str, Any]]:
    """加载 API 配置，若文件不存在则返回默认配置"""
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            logger.warning(f"配置文件读取失败，使用默认配置: {e}")
    return {
        "默认 MinerU (API)": {"api_type": "MinerU (API)", "api_key": "", "base_url": "https://mineru.net", "model_name": "vlm"},
        "默认 Local OCR": {"api_type": "Local OCR", "api_key": "", "base_url": "", "model_name": ""},
        "默认 Gemini": {"api_type": "Gemini (Native)", "api_key": "", "base_url": "", "model_name": "gemini-1.5-flash"}
    }

def save_configs(configs: Dict[str, Any]) -> None:
    """保存配置到文件"""
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(configs, f, ensure_ascii=False, indent=4)

def add_to_history(task_id: str, url_or_name: str, task_type: str, output_dir: str):
    """Add a successful task to history records"""
    history_file = os.path.join(DATA_DIR, "task_history.json")
    history_list = []
    if os.path.exists(history_file):
        try:
            with open(history_file, "r", encoding="utf-8") as f:
                history_list = json.load(f)
        except Exception:
            pass
            
    # Avoid duplicate records for the same task_id
    if any(item.get("task_id") == task_id for item in history_list):
        return
        
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    history_list.append({
        "task_id": task_id,
        "name": url_or_name,
        "type": task_type, # "url" or "upload"
        "output_dir": output_dir,
        "timestamp": timestamp
    })
    
    # Keep up to 100 history records
    if len(history_list) > 100:
        history_list = history_list[-100:]
        
    try:
        with open(history_file, "w", encoding="utf-8") as f:
            json.dump(history_list, f, ensure_ascii=False, indent=4)
    except Exception as e:
        logger.error(f"Failed to save task history: {e}")

def load_history() -> Dict[str, str]:
    """加载历史记录"""
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            logger.warning(f"历史记录读取失败: {e}")
    return {}

def save_history(history: Dict[str, str]) -> None:
    """保存历史记录"""
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=4)

def strip_code_fence(text: str) -> str:
    """去除 Markdown 代码围栏"""
    text = text.strip()
    m = re.match(r"^```[a-zA-Z]*\n?(.*?)```$", text, re.DOTALL)
    if m:
        return m.group(1).strip()
    return text

def has_markdown_table(text: str) -> bool:
    """检测文本是否包含 Markdown 或 HTML 表格"""
    text = strip_code_fence(text)
    if '<table' in text.lower() and '</table>' in text.lower():
        return True
    lines = [l.strip() for l in text.split("\n") if "|" in l]
    if len(lines) < 2:
        return False
    sep_lines = [l for l in lines if re.search(r"\|[\s]*[-:]+[\s]*\|", l)]
    return len(sep_lines) >= 1

def align_missing_name_column(cells: List[str]) -> List[str]:
    """智能对齐：如果表格行为 6 列且第一列为 4 位代号、第二列直接为科目组合，则自动插入空名称"""
    if len(cells) == 6:
        if re.match(r"^\d{4}$", cells[0]):
            if any(k in cells[1] for k in ["科目组合", "物理", "历史"]):
                return [cells[0], ""] + cells[1:]
    return cells

def parse_html_table(html_str: str) -> Optional[pd.DataFrame]:
    """手动解析 HTML 表格（不依赖 lxml/html5lib）"""
    table_match = re.search(r'<table[^>]*>(.*?)</table>', html_str, re.IGNORECASE | re.DOTALL)
    if not table_match:
        return None
    table_content = table_match.group(1)

    rows = []
    for tr_match in re.finditer(r'<tr[^>]*>(.*?)</tr>', table_content, re.IGNORECASE | re.DOTALL):
        tr_content = tr_match.group(1)
        cells = []
        for cell_match in re.finditer(r'<(td|th)[^>]*>(.*?)</\1>', tr_content, re.IGNORECASE | re.DOTALL):
            cell_text = re.sub(r'<[^>]+>', ' ', cell_match.group(2))
            cell_text = cell_text.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&').strip()
            cells.append(cell_text)
        if cells:
            m = re.match(r"^(\d{4})([一-龥]+.*)$", cells[0])
            if m:
                cells = [m.group(1), m.group(2)] + cells[1:]
            cells = align_missing_name_column(cells)
            rows.append(cells)

    if not rows:
        return None

    max_cols = max(len(r) for r in rows)
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    if len(rows) >= 2:
        return pd.DataFrame(rows[1:], columns=rows[0])
    return pd.DataFrame(rows)

def is_university_code_start(line: str, next_line: str = "") -> bool:
    """判断行是否以高校代码开头"""
    line = line.strip()
    if re.match(r"^\d{4}[一-龥]+", line):
        return True
    if re.match(r"^\d{4}$", line):
        if next_line and not next_line.isdigit() and any('一' <= c <= '龥' for c in next_line):
            return True
    return False

def reconstruct_one_cell_per_line_table(text: str) -> str:
    """修复"一词一行"格式的损坏表格数据"""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if len(lines) < 10:
        return text

    code_starts = []
    for i in range(len(lines)):
        next_line = lines[i + 1] if i + 1 < len(lines) else ""
        if is_university_code_start(lines[i], next_line):
            code_starts.append(i)

    if len(code_starts) >= 2:
        rows = []
        for idx in range(len(code_starts)):
            start = code_starts[idx]
            end = code_starts[idx + 1] if idx + 1 < len(code_starts) else len(lines)
            row_cells = lines[start:end]
            if row_cells:
                m = re.match(r"^(\d{4})([一-龥].*)$", row_cells[0])
                if m:
                    row_cells = [m.group(1), m.group(2)] + row_cells[1:]
                rows.append(row_cells)

        md_lines = []
        col_count = max(len(r) for r in rows)
        for r in rows:
            padded_row = r + [""] * (col_count - len(r))
            md_lines.append("| " + " | ".join(padded_row) + " |")
        if md_lines:
            separator = "|" + "|".join(["---"] * col_count) + "|"
            md_lines.insert(1, separator)
            return "\n".join(md_lines)

    return text

def ensure_markdown_table(text: str) -> str:
    """确保内容为标准 Markdown 表格格式"""
    text = strip_code_fence(text).strip()
    if "|" in text:
        return text
    if '<table' in text.lower() and '</table>' in text.lower():
        return text

    rebuilt_text = reconstruct_one_cell_per_line_table(text)
    if "|" in rebuilt_text:
        return rebuilt_text

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if not lines:
        return text

    if "\t" in lines[0]:
        sep = "\t"
    elif "," in lines[0]:
        sep = ","
    else:
        has_multispace = any("  " in l for l in lines)
        if has_multispace:
            rows = [re.split(r'\s{2,}', l) for l in lines]
        else:
            rows = [re.split(r'\s+', l) for l in lines]

        md_lines = []
        col_count = 0
        for row in rows:
            row_cells = [c.strip() for c in row if c.strip()]
            if row_cells:
                col_count = max(col_count, len(row_cells))
                md_lines.append("| " + " | ".join(row_cells) + " |")
        if md_lines:
            separator = "|" + "|".join(["---"] * col_count) + "|"
            md_lines.insert(1, separator)
            return "\n".join(md_lines)
        return text

    rows = [[c.strip() for c in l.split(sep) if c.strip()] for l in lines]
    md_lines = []
    col_count = 0
    for row in rows:
        if row:
            col_count = max(col_count, len(row))
            md_lines.append("| " + " | ".join(row) + " |")
    if md_lines:
        separator = "|" + "|".join(["---"] * col_count) + "|"
        md_lines.insert(1, separator)
        return "\n".join(md_lines)
    return text

def is_table_title(cells: List[str]) -> bool:
    """判定某行是否为表格标题"""
    non_empty = [c.strip() for c in cells if c.strip()]
    if len(non_empty) == 1:
        val = non_empty[0]
        if len(val) >= 10 or any(kw in val for kw in ["投档最低分", "投档最低分及名次", "录取", "招生", "投档线"]):
            return True
    return False

def extract_and_remove_table_title(md_text: str) -> tuple:
    """提取 Markdown 表格中的独立标题并剥离"""
    lines = md_text.split("\n")
    cleaned_lines = []
    title = None
    skip_next_if_separator = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append(line)
            continue

        if "|" in line:
            if skip_next_if_separator and re.match(r"^[\s\-:|]+$", stripped.replace("|", "")):
                skip_next_if_separator = False
                continue

            cells = [c.strip() for c in line.split("|")]
            if cells and cells[0] == "":
                cells = cells[1:]
            if cells and cells[-1] == "":
                cells = cells[:-1]

            if is_table_title(cells):
                title = [c.strip() for c in cells if c.strip()][0]
                skip_next_if_separator = True
                continue

        skip_next_if_separator = False
        cleaned_lines.append(line)

    return "\n".join(cleaned_lines), title

def parse_markdown_table(md_text: str) -> Optional[pd.DataFrame]:
    """解析 Markdown 表格为 DataFrame"""
    md_text = strip_code_fence(md_text)
    md_text, _ = extract_and_remove_table_title(md_text)

    if '<table' in md_text.lower() and '</table>' in md_text.lower():
        try:
            dfs = pd.read_html(io.StringIO(md_text))
            if dfs:
                return dfs[0]
        except Exception:
            df = parse_html_table(md_text)
            if df is not None:
                return df

    lines = [l.strip() for l in md_text.strip().split("\n") if "|" in l]
    data_lines = [l for l in lines if len(l.replace("|", "").replace("-", "").replace(":", "").replace(" ", "").strip()) > 0]
    if not data_lines:
        return None
    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]

        if cells and len(cells) > 0:
            m = re.match(r"^(\d{4})([一-龥]+.*)$", cells[0])
            if m:
                cells = [m.group(1), m.group(2)] + cells[1:]

        cells = align_missing_name_column(cells)
        rows.append(cells)
    if not rows:
        return None
    max_cols = max(len(r) for r in rows)
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    if len(rows) >= 2:
        return pd.DataFrame(rows[1:], columns=rows[0])
    return pd.DataFrame(rows)

def clean_control_characters(df: pd.DataFrame) -> pd.DataFrame:
    """清除 DataFrame 中的控制字符"""
    def clean_val(val):
        if isinstance(val, str):
            return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', val)
        return val
    return df.map(clean_val)

def deduplicate_columns(df: pd.DataFrame) -> pd.DataFrame:
    """对重复列名添加后缀"""
    if df.columns.is_unique:
        return df
    counts: Dict[str, int] = {}
    new_cols = []
    for col in df.columns:
        col_str = str(col)
        if col_str in counts:
            counts[col_str] += 1
            new_cols.append(f"{col_str}_{counts[col_str]}")
        else:
            counts[col_str] = 0
            new_cols.append(col_str)
    df.columns = new_cols
    return df

# ── 编译任务输出 ─────────────────────────────────────────────────────────────

def compile_task_outputs(task_id: str) -> None:
    """合并所有切片结果，生成多格式输出"""
    task = tasks_db[task_id]
    total_images = len(task["image_paths"])
    if total_images == 0:
        return

    saved_url = task.get("url", "")
    final_md = f"# 提取结果 - {saved_url}\n\n"
    final_txt = f"提取结果 - {saved_url}\n\n"

    for i in range(total_images):
        if i in task["results"]:
            img_p, res_data = task["results"][i]
            rtype = res_data.get("type", "text")
            if rtype == "ignore":
                continue
                
            content = res_data.get("content", "")
            rel_img_p = img_p.replace("\\", "/")
            if rtype == "text":
                final_md += f"*(文本 - 图 {i + 1})*\n\n{content}\n\n---\n"
                final_txt += f"{content}\n\n"
            elif rtype == "table":
                clean, title = extract_and_remove_table_title(content)
                clean = strip_code_fence(clean)
                if title:
                    final_md += f"### {title}\n\n"
                    final_txt += f"{title}\n\n"
                final_md += f"*(表格 - 图 {i + 1})*\n\n{clean}\n\n---\n"
                final_txt += f"{clean}\n\n"
            elif rtype == "error":
                final_md += f"*(错误 - 图 {i + 1})*\n\n{content}\n\n---\n"
            else:
                final_md += f"*(原图 - 图 {i + 1})*\n\n![Image]({rel_img_p})\n\n---\n"

    task["final_md"] = final_md
    task["final_txt"] = final_txt

    # 合并 Excel / CSV
    all_dfs: List[pd.DataFrame] = []
    for i in range(total_images):
        if i in task["results"]:
            _, res_data = task["results"][i]
            rtype = res_data.get("type", "text")
            if rtype == "ignore":
                continue
            content = res_data.get("content", "")
            if rtype == "table" or has_markdown_table(content):
                try:
                    table_md = ensure_markdown_table(content)
                    df = parse_markdown_table(table_md)
                    if df is not None and not df.empty:
                        df = deduplicate_columns(df)
                        all_dfs.append(df)
                except Exception:
                    pass

    if all_dfs:
        try:
            merged_df = _align_and_merge_dfs(all_dfs)
        except Exception as e:
            logger.error(f"表格列对齐合并算法运行出错，降级回退至 Pandas 标准拼合: {e}")
            merged_df = pd.concat(all_dfs, ignore_index=True)

        merged_df = clean_control_characters(merged_df)

        csv_buf = io.StringIO()
        merged_df.to_csv(csv_buf, index=False)
        task["csv_bytes"] = csv_buf.getvalue().encode("utf-8-sig")

        excel_buf = io.BytesIO()
        try:
            save_df_to_excel_with_merged_cells(merged_df, excel_buf, sheet_name="合并表格")
        except Exception as excel_err:
            logger.error(f"合并单元格导出 Excel 失败，退回到 Pandas 默认导出: {excel_err}")
            excel_buf = io.BytesIO()
            merged_df.to_excel(excel_buf, index=False, engine="openpyxl", sheet_name="合并表格")
        excel_buf.seek(0)
        task["excel_bytes"] = excel_buf.getvalue()

    # 保存结果到本地文件
    output_dir = task["output_dir"]
    if output_dir and os.path.exists(output_dir):
        _save_outputs_to_disk(task, output_dir)


def _align_and_merge_dfs(dfs: List[pd.DataFrame]) -> pd.DataFrame:
    """智能对齐列结构合并算法"""
    base_df = dfs[0].copy()
    base_cols = [str(c).strip() for c in base_df.columns]
    base_df.columns = base_cols

    aligned_dfs = [base_df]

    def is_duplicate_header(cols1: List[str], cols2: List[str]) -> bool:
        matches = sum(1 for c1, c2 in zip(cols1, cols2) if c1.lower() == c2.lower())
        return matches >= max(1, len(cols2) // 2)

    for df_orig in dfs[1:]:
        df = df_orig.copy()
        df_cols = [str(c).strip() for c in df.columns]
        df.columns = df_cols

        if len(df_cols) == len(base_cols):
            if is_duplicate_header(df_cols, base_cols):
                df.columns = base_cols
                aligned_dfs.append(df)
            else:
                header_row = list(df.columns)
                data_list = [header_row] + df.values.tolist()
                new_df = pd.DataFrame(data_list, columns=base_cols)
                aligned_dfs.append(new_df)
        elif len(df_cols) < len(base_cols):
            if is_duplicate_header(df_cols, base_cols[:len(df_cols)]):
                df.columns = base_cols[:len(df_cols)]
                for extra_col in base_cols[len(df_cols):]:
                    df[extra_col] = ""
                aligned_dfs.append(df[base_cols])
            else:
                header_row = list(df.columns) + [""] * (len(base_cols) - len(df_cols))
                data_list = [list(row) + [""] * (len(base_cols) - len(row)) for row in df.values.tolist()]
                data_list = [header_row] + data_list
                new_df = pd.DataFrame(data_list, columns=base_cols)
                aligned_dfs.append(new_df)
        else:
            if is_duplicate_header(df_cols[:len(base_cols)], base_cols):
                sub_df = df.iloc[:, :len(base_cols)].copy()
                sub_df.columns = base_cols
                aligned_dfs.append(sub_df)
            else:
                header_row = list(df.columns)[:len(base_cols)]
                data_list = [list(row)[:len(base_cols)] for row in df.values.tolist()]
                data_list = [header_row] + data_list
                new_df = pd.DataFrame(data_list, columns=base_cols)
                aligned_dfs.append(new_df)

    return pd.concat(aligned_dfs, ignore_index=True)


def save_df_to_excel_with_merged_cells(df: pd.DataFrame, file_or_buf, sheet_name: str = "合并表格") -> None:
    """利用 openpyxl 支持合并相邻相同单元格（垂直/水平）导出到 Excel 的增强导出函数"""
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Side, PatternFill, Font
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    # 显式开启网格线显示，防止部分 Excel 软件因渲染模式导致边框不可见
    ws.views.sheetView[0].showGridLines = True

    # 1. 写入表头和所有数据
    headers = [str(c) for c in df.columns]
    ws.append(headers)
    for row in df.values:
        ws.append([("" if pd.isna(v) else str(v)) for v in row])

    num_rows = ws.max_row
    num_cols = ws.max_column

    # 样式参数定义
    thin_border = Border(
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='thin', color='D3D3D3'),
        top=Side(style='thin', color='D3D3D3'),
        bottom=Side(style='thin', color='D3D3D3')
    )
    header_fill = PatternFill(start_color="3B82F6", end_color="3B82F6", fill_type="solid")
    header_font = Font(name="Microsoft YaHei", size=10, bold=True, color="FFFFFF")
    data_font = Font(name="Microsoft YaHei", size=10)

    # 2. 设置表头格式
    for col in range(1, num_cols + 1):
        cell = ws.cell(row=1, column=col)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border

    # 3. 垂直合并相邻且相同的单元格已关闭（根据用户要求，一般只做水平合并）

    def is_numeric(v) -> bool:
        if v is None:
            return False
        s = str(v).strip()
        try:
            float(s)
            return True
        except ValueError:
            return False

    # 4. 合并每一行中相邻且相同的水平单元格 (数据行水平合并，跳过数字/分数以防错合并)
    for row in range(2, num_rows + 1):
        start_col = 1
        while start_col <= num_cols:
            val = ws.cell(row=row, column=start_col).value
            if val is None or str(val).strip() == "" or is_numeric(val):
                start_col += 1
                continue

            end_col = start_col
            while end_col + 1 <= num_cols and ws.cell(row=row, column=end_col + 1).value == val:
                end_col += 1

            if end_col > start_col:
                try:
                    ws.merge_cells(start_row=row, start_column=start_col, end_row=row, end_column=end_col)
                    ws.cell(row=row, column=start_col).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                except Exception:
                    pass
                start_col = end_col + 1
            else:
                start_col += 1

    # 5. 格式化所有单元格字体、对齐和边框，解决合并单元格缺失边框的问题
    for r in range(2, num_rows + 1):
        for c in range(1, num_cols + 1):
            cell = ws.cell(row=r, column=c)
            cell.font = data_font
            cell.border = thin_border
            if not cell.alignment:
                cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # 6. 自适应列宽自动调整
    for col in range(1, num_cols + 1):
        max_len = 10
        for r in range(1, num_rows + 1):
            val = ws.cell(row=r, column=col).value
            if val is not None:
                val_str = str(val)
                # 区分中英文长度计算
                length = sum(2 if ord(char) > 127 else 1 for char in val_str)
                if length > max_len:
                    max_len = length
        col_letter = get_column_letter(col)
        ws.column_dimensions[col_letter].width = min(max_len + 3, 50)

    wb.save(file_or_buf)


def _save_outputs_to_disk(task: Dict[str, Any], output_dir: str) -> None:
    """将编译后的输出保存到磁盘"""
    try:
        with open(os.path.join(output_dir, "final_output.md"), "w", encoding="utf-8") as f:
            f.write(task["final_md"])
        with open(os.path.join(output_dir, "final_output.txt"), "w", encoding="utf-8") as f:
            f.write(task["final_txt"])
        if task.get("csv_bytes"):
            with open(os.path.join(output_dir, "final_output.csv"), "wb") as f:
                f.write(task["csv_bytes"])
        if task.get("excel_bytes"):
            with open(os.path.join(output_dir, "final_output.xlsx"), "wb") as f:
                f.write(task["excel_bytes"])

        html_content = md_to_html(task["final_md"])
        with open(os.path.join(output_dir, "final_output.html"), "w", encoding="utf-8") as f:
            f.write(html_content)

        latex_content = md_to_latex(task["final_md"])
        with open(os.path.join(output_dir, "final_output.tex"), "w", encoding="utf-8") as f:
            f.write(latex_content)

        docx_bytes = md_to_docx(task["final_md"])
        with open(os.path.join(output_dir, "final_output.docx"), "wb") as f:
            f.write(docx_bytes)
    except Exception as e:
        logger.warning(f"无法写入最终本地合并文件: {e}")
        task["logs"].append(
            f"[{datetime.datetime.now().strftime('%H:%M:%S')}] 警告: 无法写入最终合并文件: {e}"
        )


# ── 任务处理器核心逻辑 ──────────────────────────────────────────────────────

def run_extraction_task(task_id: str, url: str, config: Dict[str, Any], loop: asyncio.AbstractEventLoop) -> None:
    """后台线程执行网页提取任务"""
    task = tasks_db[task_id]

    def log_progress(msg: str) -> None:
        logger.info(f"[{task_id}] {msg}")
        task["logs"].append(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] {msg}")
        if task_id in tasks_events:
            loop.call_soon_threadsafe(tasks_events[task_id].set)

    try:
        log_progress("开始处理网页链接...")

        history = load_history()
        output_dir = history.get(url)
        image_paths_file = os.path.join(output_dir, "image_paths.json") if output_dir else ""

        breakpoint_hit = False
        image_paths: List[str] = []

        # 尝试断点续传
        if output_dir and os.path.exists(output_dir) and os.path.exists(image_paths_file):
            log_progress("检测到已存在历史抓取记录，正在加载断点...")
            try:
                with open(image_paths_file, "r", encoding="utf-8") as f:
                    image_paths = json.load(f)

                task["output_dir"] = output_dir
                task["image_paths"] = image_paths
                task["total_count"] = len(image_paths)
                breakpoint_hit = True

                for idx, img_path in enumerate(image_paths):
                    slice_md = os.path.join(output_dir, f"slice_{idx + 1}.md")
                    slice_csv = os.path.join(output_dir, f"table_{idx + 1}.csv")

                    if os.path.exists(slice_md):
                        with open(slice_md, "r", encoding="utf-8") as f:
                            content = f.read()

                        content = ensure_markdown_table(content)
                        rtype = "image" if content == "RETAIN_IMAGE" else ("table" if has_markdown_table(content) else "text")
                        res_data: Dict[str, Any] = {"type": rtype, "content": content}
                        if os.path.exists(slice_csv):
                            res_data["csv_path"] = slice_csv

                        with task["lock"]:
                            task["results"][idx] = (img_path, res_data)
                            task["completed_count"] += 1

                log_progress(f"断点加载完成！已跳过已处理的 {task['completed_count']}/{task['total_count']} 张切片。")
            except Exception as e:
                log_progress(f"加载断点失败，将重新开始: {e}")
                output_dir = None
                task["results"] = {}
                task["completed_count"] = 0

        # 全新任务流程
        if not output_dir or not os.path.exists(image_paths_file):
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = os.path.join(OUTPUT_DIR, timestamp)
            os.makedirs(output_dir, exist_ok=True)
            task["output_dir"] = output_dir

            history[url] = output_dir
            save_history(history)

            log_progress("正在初始化抓取网页并提取大图...")
            scraper = WebScraper(output_dir=output_dir)
            original_image_paths = scraper.fetch_images(url)

            if not original_image_paths:
                log_progress("警告: 未在网页中找到符合条件的图片。")
                task["status"] = "completed"
                log_progress("任务结束（未抓取到有效图片）。")
                return

            log_progress(f"成功抓取到 {len(original_image_paths)} 张原图。正在进行智能无缝切分...")
            for img_path in original_image_paths:
                slices = smart_slice(img_path, max_height=1000)
                image_paths.extend(slices)

            task["image_paths"] = image_paths
            task["total_count"] = len(image_paths)

            with open(os.path.join(output_dir, "image_paths.json"), "w", encoding="utf-8") as f:
                json.dump(image_paths, f, ensure_ascii=False, indent=4)

            log_progress(f"长图智能切分完成！共生成 {len(image_paths)} 张切片。")

        # 启动识别引擎并发处理剩余图片
        total_images = len(image_paths)
        unprocessed_indices = [i for i in range(total_images) if i not in task["results"]]

        if unprocessed_indices:
            log_progress(f"正在配置 {config['api_type']} 识别引擎，启用并行处理...")
            analyzer = ImageAnalyzer(
                api_type=config["api_type"],
                api_key=config["api_key"],
                base_url=config["base_url"],
                model_name=config["model_name"],
                split_len=config.get("split_len"),
                split_regex=config.get("split_regex")
            )

            _image_paths = list(image_paths)
            _output_dir = output_dir

            def process_single_image(idx: int):
                while task["status"] == "paused":
                    time.sleep(0.5)
                if task["status"] == "stopped":
                    raise Exception("Task stopped by user")

                path = _image_paths[idx]
                out_csv = os.path.join(_output_dir, f"table_{idx + 1}.csv")
                res = analyzer.analyze(path, output_csv_path=out_csv, formats=task.get("formats"))

                if task["status"] == "stopped":
                    raise Exception("Task stopped by user")

                out_md = os.path.join(_output_dir, f"slice_{idx + 1}.md")
                with open(out_md, "w", encoding="utf-8") as f:
                    if res["type"] in ["text", "table"]:
                        f.write(res["content"])
                    elif res["type"] == "image":
                        f.write("RETAIN_IMAGE")
                    else:
                        f.write(res["content"])
                return idx, path, res

            with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
                future_to_idx = {executor.submit(process_single_image, i): i for i in unprocessed_indices}
                task["current_processing"] = set(unprocessed_indices)
                if task_id in tasks_events:
                    loop.call_soon_threadsafe(tasks_events[task_id].set)

                for future in concurrent.futures.as_completed(future_to_idx):
                    if task["status"] == "stopped":
                        break
                    idx = future_to_idx[future]
                    task["current_processing"].discard(idx)
                    try:
                        _, img_p, res_data = future.result()
                        if res_data["type"] != "image" and has_markdown_table(res_data.get("content", "")):
                            res_data["type"] = "table"
                        with task["lock"]:
                            task["results"][idx] = (img_p, res_data)
                        log_progress(f"图片 {idx + 1}/{total_images} 识别完成。")
                    except Exception as exc:
                        if "Task stopped by user" in str(exc) or task["status"] == "stopped":
                            continue
                        log_progress(f"图片 {idx + 1} 识别失败: {exc}")
                        with task["lock"]:
                            task["results"][idx] = (_image_paths[idx], {"type": "error", "content": str(exc)})

                    task["completed_count"] += 1
                    task["progress"] = int((task["completed_count"] / total_images) * 100)
                    compile_task_outputs(task_id)

        if task["status"] == "stopped":
            log_progress("任务已被用户停止。")
            return

        log_progress("正在进行多格式文档合并编译...")
        compile_task_outputs(task_id)
        log_progress("恭喜！网页多模态提取任务全部完成，报告已生成。")
        task["progress"] = 100
        task["status"] = "completed"
        add_to_history(task_id, url, "url", output_dir)

    except Exception as e:
        task["status"] = "failed"
        task["error"] = str(e)
        log_progress(f"任务执行中遭遇严重错误崩溃: {e}")
        logger.exception(e)


def run_upload_task(
    task_id: str, file_paths: List[str],
    config: Dict[str, Any], loop: asyncio.AbstractEventLoop, output_dir: str,
    url: str = ""
) -> None:
    """处理上传文件的任务（支持单/多文件混合，支持与网页 URL 合并）"""
    task = tasks_db[task_id]

    def log_progress(msg: str) -> None:
        logger.info(f"[{task_id}] {msg}")
        task["logs"].append(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] {msg}")
        if task_id in tasks_events:
            loop.call_soon_threadsafe(tasks_events[task_id].set)

    try:
        log_progress("开始处理提取队列项目...")
        image_exts = [".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"]
        task_items = []

        # 1. 如果有网页 URL，先抓取网页图片并分片
        if url:
            log_progress(f"正在抓取并解析网页链接: {url} ...")
            scraper = WebScraper(output_dir=output_dir)
            original_image_paths = scraper.fetch_images(url)
            if original_image_paths:
                log_progress(f"成功抓取到 {len(original_image_paths)} 张网页原图，正在进行切分排队...")
                for img_p in original_image_paths:
                    slices = smart_slice(img_p, max_height=1000)
                    task_items.extend(slices)
            else:
                log_progress("警告: 未在网页中抓取到任何符合条件的图片。")

        # 2. 处理本地上传的文件队列
        if file_paths:
            log_progress(f"开始处理已上传的 {len(file_paths)} 个本地文件...")
            for p in file_paths:
                filename = os.path.basename(p)
                _, ext = os.path.splitext(p)
                ext = ext.lower()
                if ext in image_exts:
                    log_progress(f"正在对图片 {filename} 进行智能无缝切分...")
                    slices = smart_slice(p, max_height=1000)
                    task_items.extend(slices)
                else:
                    log_progress(f"本地文档 {filename} 已成功加入待处理队列。")
                    task_items.append(p)

        if not task_items:
            raise ValueError("没有需要解析的有效图片或文档对象")

        task["image_paths"] = task_items
        task["total_count"] = len(task_items)

        with open(os.path.join(output_dir, "image_paths.json"), "w", encoding="utf-8") as f:
            json.dump(task_items, f, ensure_ascii=False, indent=4)

        log_progress(f"文件处理与切片生成完毕，共 {len(task_items)} 个待提取项。配置 {config['api_type']} 引擎中...")

        analyzer = ImageAnalyzer(
            api_type=config["api_type"],
            api_key=config["api_key"],
            base_url=config["base_url"],
            model_name=config["model_name"],
            split_len=config.get("split_len"),
            split_regex=config.get("split_regex")
        )

        _task_items = list(task_items)
        _output_dir = output_dir

        def process_single(idx: int):
            while task["status"] == "paused":
                time.sleep(0.5)
            if task["status"] == "stopped":
                raise Exception("Task stopped by user")

            path = _task_items[idx]
            out_csv = os.path.join(_output_dir, f"table_{idx + 1}.csv")
            res = analyzer.analyze(path, output_csv_path=out_csv, formats=task.get("formats"))

            if task["status"] == "stopped":
                raise Exception("Task stopped by user")

            out_md = os.path.join(_output_dir, f"slice_{idx + 1}.md")
            with open(out_md, "w", encoding="utf-8") as f:
                f.write(res["content"] if res["type"] in ["text", "table"] else "RETAIN_IMAGE")
            return idx, path, res

        total = len(task_items)
        unprocessed_indices = [i for i in range(total) if i not in task["results"]]
        task["current_processing"] = set(unprocessed_indices)
        if task_id in tasks_events:
            loop.call_soon_threadsafe(tasks_events[task_id].set)

        if unprocessed_indices:
            with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
                futures = {executor.submit(process_single, i): i for i in unprocessed_indices}
                for future in concurrent.futures.as_completed(futures):
                    if task["status"] == "stopped":
                        break
                    idx = futures[future]
                    task["current_processing"].discard(idx)
                    try:
                        _, img_p, res_data = future.result()
                        if res_data["type"] != "image" and has_markdown_table(res_data.get("content", "")):
                            res_data["type"] = "table"
                        with task["lock"]:
                            task["results"][idx] = (img_p, res_data)
                        log_progress(f"提取项 {idx + 1}/{total} 识别完成。")
                    except Exception as exc:
                        if "Task stopped by user" in str(exc) or task["status"] == "stopped":
                            continue
                        log_progress(f"提取项 {idx + 1} 识别失败: {exc}")
                        with task["lock"]:
                            task["results"][idx] = (_task_items[idx], {"type": "error", "content": str(exc)})
                    task["completed_count"] += 1
                    task["progress"] = int((task["completed_count"] / total) * 100)
                    compile_task_outputs(task_id)

        if task["status"] == "stopped":
            log_progress("任务已被用户停止。")
            return

        compile_task_outputs(task_id)
        task["progress"] = 100
        task["status"] = "completed"
        log_progress("队列中所有对象提取与报告合并编译全部完成！")
        name = url if url else ", ".join([os.path.basename(p) for p in file_paths])
        add_to_history(task_id, name, "url" if url else "upload", output_dir)

    except Exception as e:
        task["status"] = "failed"
        task["error"] = str(e)
        log_progress(f"任务提取线程遭遇严重错误崩溃: {e}")
        logger.exception(e)




# ── 格式转换工具 ─────────────────────────────────────────────────────────────

def md_to_html(md_text: str) -> str:
    """将 Markdown 转为简易 HTML"""
    import html as html_lib
    lines = md_text.split("\n")
    html_lines: List[str] = []
    in_table = False
    in_code = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                html_lines.append("</pre>")
                in_code = False
            else:
                html_lines.append("<pre style='background:#f5f5f5;padding:1em;border-radius:8px;overflow-x:auto;'>")
                in_code = True
            continue

        if in_code:
            html_lines.append(html_lib.escape(line))
            continue

        if stripped.startswith("<") and not stripped.startswith("<br"):
            html_lines.append(line)
            continue

        if stripped.startswith("# "):
            html_lines.append(f"<h1>{stripped[2:]}</h1>")
        elif stripped.startswith("## "):
            html_lines.append(f"<h2>{stripped[3:]}</h2>")
        elif stripped.startswith("### "):
            html_lines.append(f"<h3>{stripped[4:]}</h3>")
        elif stripped == "---":
            html_lines.append("<hr>")
        elif stripped.startswith("|") and "|" in stripped:
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if re.match(r"^[\s\-:|]+$", stripped.replace("|", "")):
                continue
            if not in_table:
                html_lines.append("<table border='1' style='border-collapse:collapse;width:100%;'>")
                html_lines.append(
                    "<tr>" + "".join(f"<th style='padding:6px;border:1px solid #ddd;'>{c}</th>" for c in cells) + "</tr>"
                )
                in_table = True
            else:
                html_lines.append(
                    "<tr>" + "".join(f"<td style='padding:6px;border:1px solid #ddd;'>{c}</td>" for c in cells) + "</tr>"
                )
        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            html_lines.append(f"<p><em>{stripped[1:-1]}</em></p>")
        elif stripped == "":
            if in_table:
                html_lines.append("</table>")
                in_table = False
            html_lines.append("")
        else:
            processed = stripped.replace("**", "<strong>", 1).replace("**", "</strong>", 1)
            html_lines.append(f"<p>{processed}</p>")

    if in_table:
        html_lines.append("</table>")

    body = "\n".join(html_lines)
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head><meta charset="UTF-8"><title>提取结果</title>
<style>body{{font-family:sans-serif;max-width:900px;margin:2rem auto;padding:0 1rem;line-height:1.6;}}
table{{margin:1rem 0;}} th{{background:#f0f0f0;}} img{{max-width:100%;}}</style>
</head><body>{body}</body></html>"""


def md_to_latex(md_text: str) -> str:
    """将 Markdown 转为简易 LaTeX"""
    lines = md_text.split("\n")
    tex_lines: List[str] = [
        r"\documentclass[a4paper,12pt]{article}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage{longtable}",
        r"\usepackage{graphicx}",
        r"\usepackage[margin=2cm]{geometry}",
        r"\begin{document}",
    ]
    in_table = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            tex_lines.append(r"\section*{" + _latex_escape(stripped[2:]) + "}")
        elif stripped.startswith("## "):
            tex_lines.append(r"\subsection*{" + _latex_escape(stripped[3:]) + "}")
        elif stripped.startswith("### "):
            tex_lines.append(r"\subsubsection*{" + _latex_escape(stripped[4:]) + "}")
        elif stripped == "---":
            tex_lines.append(r"\hrule\vspace{0.5em}")
        elif stripped.startswith("|") and "|" in stripped:
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if re.match(r"^[\s\-:|]+$", stripped.replace("|", "")):
                continue
            if not in_table:
                num_cols = len(cells)
                tex_lines.append(r"\begin{longtable}{" + "|".join(["l"] * num_cols) + "}")
                tex_lines.append(r"\hline")
                tex_lines.append(" & ".join(_latex_escape(c) for c in cells) + r" \\ \hline")
                in_table = True
            else:
                tex_lines.append(" & ".join(_latex_escape(c) for c in cells) + r" \\")
        elif stripped == "" and in_table:
            tex_lines.append(r"\hline")
            tex_lines.append(r"\end{longtable}")
            in_table = False
        elif stripped.startswith("*") and stripped.endswith("*"):
            tex_lines.append(r"\textit{" + _latex_escape(stripped[1:-1]) + "}")
        elif stripped:
            cleaned = stripped.replace("<br>", r"\\").replace("<br/>", r"\\")
            tex_lines.append(_latex_escape(cleaned))
        else:
            tex_lines.append("")

    if in_table:
        tex_lines.append(r"\hline")
        tex_lines.append(r"\end{longtable}")

    tex_lines.append(r"\end{document}")
    return "\n".join(tex_lines)


def _latex_escape(text: str) -> str:
    """转义 LaTeX 特殊字符"""
    chars = {
        '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#', '_': r'\_',
        '{': r'\{', '}': r'\}', '~': r'\textasciitilde{}', '^': r'\textasciicircum{}'
    }
    for k, v in chars.items():
        text = text.replace(k, v)
    return text


def md_to_docx(md_text: str) -> bytes:
    """将 Markdown 转为 DOCX 字节流"""
    try:
        from docx import Document
    except ImportError:
        buf = io.BytesIO()
        buf.write(md_text.encode("utf-8"))
        buf.seek(0)
        return buf.getvalue()

    doc = Document()
    lines = md_text.split("\n")

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped == "---":
            doc.add_paragraph("─" * 40)
        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped[1:-1])
            run.italic = True
        elif stripped == "":
            continue
        else:
            cleaned = stripped.replace("<br>", "\n").replace("<br/>", "\n")
            doc.add_paragraph(cleaned)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── 静态文件托管 ─────────────────────────────────────────────────────────────

STATIC_DIR = os.path.join(BASE_DIR, "static")
os.makedirs(STATIC_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
app.mount("/output", StaticFiles(directory=OUTPUT_DIR), name="output")


# ── API 路由 ─────────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
def get_home() -> str:
    """返回前端页面"""
    index_path = os.path.join(BASE_DIR, "index.html")
    if os.path.exists(index_path):
        with open(index_path, "r", encoding="utf-8") as f:
            return f.read()
    return "<h3>Frontend index.html not found!</h3>"


@app.get("/api/config")
def get_config() -> Dict[str, Any]:
    """获取所有配置"""
    return {"configs": load_configs()}


@app.post("/api/config")
def update_config(data: Dict[str, Any] = Body(...)) -> Dict[str, Any]:
    """新增或更新配置"""
    name = data.get("name")
    config_data = data.get("config")
    if not name or not config_data:
        raise HTTPException(status_code=400, detail="Missing configuration details")

    configs = load_configs()
    configs[name] = config_data
    save_configs(configs)
    return {"status": "success", "configs": configs}


@app.delete("/api/config/{name}")
def delete_config(name: str) -> Dict[str, Any]:
    """删除配置"""
    configs = load_configs()
    if name in configs:
        del configs[name]
        save_configs(configs)
        return {"status": "success", "configs": configs}
    raise HTTPException(status_code=404, detail="Configuration not found")


def _create_task(source_label: str, config: Dict[str, Any], formats: Optional[Dict[str, Any]] = None) -> str:
    """创建通用任务对象"""
    task_id = str(uuid.uuid4())
    tasks_db[task_id] = {
        "url": source_label,
        "config": config,
        "formats": formats or {"text": True, "table": True},
        "status": "running",
        "progress": 0,
        "logs": [],
        "completed_count": 0,
        "total_count": 0,
        "results": {},
        "lock": threading.Lock(),
        "image_paths": [],
        "output_dir": "",
        "final_md": "",
        "final_txt": "",
        "csv_bytes": None,
        "excel_bytes": None,
        "error": "",
        "current_processing": set()
    }
    tasks_events[task_id] = asyncio.Event()
    return task_id


@app.post("/api/extract")
async def start_extract(data: Dict[str, Any] = Body(...)) -> Dict[str, Any]:
    """开始网页提取任务"""
    url = data.get("url")
    config_name = data.get("config_name")
    force = data.get("force", False)
    reuse = data.get("reuse", False)
    formats = data.get("formats")

    if not url:
        raise HTTPException(status_code=400, detail="请输入有效的网页链接")

    url = url.strip()
    if not url.startswith(("http://", "https://")):
        url = "https://" + url

    configs = load_configs()
    config = configs.get(config_name)
    if not config:
        config = list(configs.values())[0] if configs else {
            "api_type": "Local OCR", "api_key": "", "base_url": "", "model_name": ""
        }

    history = load_history()
    output_dir = history.get(url)
    has_report = False
    if output_dir and os.path.exists(output_dir):
        report_files = [
            "final_output.md", "final_output.txt", "final_output.csv",
            "final_output.xlsx", "final_output.html", "final_output.tex",
            "final_output.docx", "final_output.doc"
        ]
        if any(os.path.exists(os.path.join(output_dir, f)) for f in report_files):
            has_report = True

    if has_report and not force and not reuse:
        return {"status": "already_exists", "message": "检测到该链接的历史报告已合并生成。"}

    if has_report and force:
        try:
            report_files = [
                "final_output.md", "final_output.txt", "final_output.csv",
                "final_output.xlsx", "final_output.html", "final_output.tex",
                "final_output.docx", "final_output.doc"
            ]
            for f in report_files:
                f_path = os.path.join(output_dir, f)
                if os.path.exists(f_path):
                    os.remove(f_path)
        except Exception as e:
            logger.warning(f"删除历史最终报告失败: {e}")

    task_id = _create_task(url, config, formats)
    loop = asyncio.get_running_loop()

    if reuse and output_dir and os.path.exists(output_dir):
        task = tasks_db[task_id]
        task["output_dir"] = output_dir

        image_paths_file = os.path.join(output_dir, "image_paths.json")
        image_paths: List[str] = []
        if os.path.exists(image_paths_file):
            try:
                with open(image_paths_file, "r", encoding="utf-8") as f:
                    image_paths = json.load(f)
            except Exception:
                pass

        task["image_paths"] = image_paths
        task["total_count"] = len(image_paths)

        for idx in range(task["total_count"]):
            slice_md = os.path.join(output_dir, f"slice_{idx + 1}.md")
            if os.path.exists(slice_md):
                with open(slice_md, "r", encoding="utf-8") as f:
                    md_content = f.read()

                md_content = ensure_markdown_table(md_content)
                res_data: Dict[str, Any] = {"type": "text", "content": md_content}
                if md_content.strip() == "RETAIN_IMAGE":
                    res_data["type"] = "image"
                elif has_markdown_table(md_content):
                    res_data["type"] = "table"

                slice_csv = os.path.join(output_dir, f"table_{idx + 1}.csv")
                if os.path.exists(slice_csv):
                    res_data["csv_path"] = slice_csv

                task["results"][idx] = (image_paths[idx], res_data)
                task["completed_count"] += 1

        compile_task_outputs(task_id)
        task["status"] = "completed"
        task["progress"] = 100
        task["logs"].append(
            f"[{datetime.datetime.now().strftime('%H:%M:%S')}] [系统] 检测到已有历史报告，直接加载并跳过。"
        )
        if task_id in tasks_events:
            loop.call_soon_threadsafe(tasks_events[task_id].set)
    else:
        thread = threading.Thread(
            target=run_extraction_task,
            args=(task_id, url, config, loop),
            daemon=True
        )
        thread.start()

    return {"status": "started", "task_id": task_id}


@app.post("/api/upload")
async def upload_file(
    files: List[UploadFile] = File(...),
    config_name: str = Form(""),
    formats: Optional[str] = Form(None),
    url: Optional[str] = Form("")
) -> Dict[str, Any]:
    """上传一个或多个图片/PDF文件进行解析"""
    if not files or len(files) == 0:
        raise HTTPException(status_code=400, detail="请选择文件")

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = os.path.join(OUTPUT_DIR, f"upload_{timestamp}")
    os.makedirs(output_dir, exist_ok=True)

    saved_paths = []
    for file in files:
        if not file.filename:
            continue
        ext = os.path.splitext(file.filename)[1].lower()
        allowed = [".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp",
                   ".pdf", ".ppt", ".pptx"]
        if ext not in allowed:
            raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

        saved_path = os.path.join(output_dir, file.filename)
        with open(saved_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
        saved_paths.append(saved_path)

    if not saved_paths:
        raise HTTPException(status_code=400, detail="未上传任何有效文件")

    configs = load_configs()
    config = configs.get(config_name)
    if not config:
        config = list(configs.values())[0] if configs else {
            "api_type": "Local OCR", "api_key": "", "base_url": "", "model_name": ""
        }

    formats_dict = None
    if formats:
        try:
            formats_dict = json.loads(formats)
        except Exception:
            pass

    display_name = f"上传: {os.path.basename(saved_paths[0])}"
    if len(saved_paths) > 1:
        display_name += f" 等 {len(saved_paths)} 个文件"

    # 如果有 URL 则是合并任务，更新任务名称
    if url:
        display_name = f"合并提取: {url} + {display_name}"

    task_id = _create_task(display_name, config, formats_dict)
    tasks_db[task_id]["output_dir"] = output_dir
    tasks_db[task_id]["file_paths"] = saved_paths
    tasks_db[task_id]["file_path"] = saved_paths[0]  # 向后兼容旧字段
    tasks_db[task_id]["file_ext"] = os.path.splitext(saved_paths[0])[1].lower()
    if url:
        tasks_db[task_id]["merged_url"] = url

    loop = asyncio.get_running_loop()
    thread = threading.Thread(
        target=run_upload_task,
        args=(task_id, saved_paths, config, loop, output_dir, url),
        daemon=True
    )
    thread.start()
    return {"task_id": task_id}


@app.get("/api/events/{task_id}")
async def get_task_events(task_id: str) -> EventSourceResponse:
    """SSE 实时推送任务进度"""
    if task_id not in tasks_db:
        raise HTTPException(status_code=404, detail="Task not found")

    def build_sse_payload(task: Dict[str, Any]) -> str:
        with task["lock"]:
            results_copy = dict(task["results"])
            logs_copy = list(task["logs"])
            processing_copy = list(task.get("current_processing", set()))
            status = task["status"]
            progress = task["progress"]
            completed = task["completed_count"]
            total = task["total_count"]
            error = task["error"]

        formatted_results = {}
        for k, v in results_copy.items():
            img_path = v[0]
            res_data = v[1]
            normalized_path = img_path.replace("\\", "/")
            if "/output/" in normalized_path:
                idx = normalized_path.rindex("/output/")
                url_path = normalized_path[idx+1:]
            else:
                try:
                    rel = os.path.relpath(img_path, OUTPUT_DIR).replace("\\", "/")
                    url_path = f"output/{rel}"
                except Exception:
                    url_path = normalized_path
            formatted_results[str(k)] = [
                url_path, 
                {
                    "type": res_data["type"], 
                    "content": res_data["content"], 
                    "source": res_data.get("source", "Local OCR")
                }
            ]

        return json.dumps({
            "status": status,
            "progress": progress,
            "completed": completed,
            "total": total,
            "logs": logs_copy,
            "error": error,
            "current_processing": sorted(processing_copy),
            "results": formatted_results
        }, ensure_ascii=False)

    async def event_generator():
        task = tasks_db[task_id]
        yield {"data": build_sse_payload(task)}

        while task["status"] in ["running", "paused"]:
            try:
                await asyncio.wait_for(tasks_events[task_id].wait(), timeout=1.0)
                tasks_events[task_id].clear()
            except asyncio.TimeoutError:
                pass
            yield {"data": build_sse_payload(task)}

        yield {"data": build_sse_payload(task)}

    return EventSourceResponse(event_generator())


@app.post("/api/retry/{task_id}")
async def retry_task(task_id: str) -> Dict[str, Any]:
    """从故障中重新启动任务"""
    if task_id not in tasks_db:
        raise HTTPException(status_code=404, detail="Task not found")
    task = tasks_db[task_id]
    if task["status"] not in ["failed", "running"]:
        raise HTTPException(status_code=400, detail="Only failed or halted tasks can be retried")

    task["status"] = "running"
    task["error"] = ""

    loop = asyncio.get_running_loop()
    config = task.get("config", {"api_type": "Local OCR", "api_key": "", "base_url": "", "model_name": ""})

    url = task.get("url", "")
    if (url.startswith("上传: ") or url.startswith("合并提取: ")) and (task.get("file_paths") or task.get("file_path")):
        file_paths = task.get("file_paths") or [task["file_path"]]
        output_dir = task.get("output_dir")
        merged_url = task.get("merged_url", "")
        thread = threading.Thread(
            target=run_upload_task,
            args=(task_id, file_paths, config, loop, output_dir, merged_url),
            daemon=True
        )
    else:
        thread = threading.Thread(
            target=run_extraction_task,
            args=(task_id, url, config, loop),
            daemon=True
        )

    thread.start()
    return {"status": "success"}


@app.post("/api/pause/{task_id}")
async def pause_task(task_id: str) -> Dict[str, Any]:
    """暂停运行中的任务"""
    if task_id not in tasks_db:
        raise HTTPException(status_code=404, detail="Task not found")
    task = tasks_db[task_id]
    if task["status"] != "running":
        raise HTTPException(status_code=400, detail="Only running tasks can be paused")
    task["status"] = "paused"
    task["logs"].append(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] 任务已被用户暂停。")
    if task_id in tasks_events:
        tasks_events[task_id].set()
    return {"status": "paused"}


@app.post("/api/resume/{task_id}")
async def resume_task(task_id: str) -> Dict[str, Any]:
    """继续已暂停的任务"""
    if task_id not in tasks_db:
        raise HTTPException(status_code=404, detail="Task not found")
    task = tasks_db[task_id]
    if task["status"] != "paused":
        raise HTTPException(status_code=400, detail="Only paused tasks can be resumed")
    task["status"] = "running"
    task["logs"].append(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] 任务已被用户继续恢复运行。")
    if task_id in tasks_events:
        tasks_events[task_id].set()
    return {"status": "running"}


@app.post("/api/stop/{task_id}")
async def stop_task(task_id: str) -> Dict[str, Any]:
    """终止/停止运行中或暂停的任务"""
    if task_id not in tasks_db:
        raise HTTPException(status_code=404, detail="Task not found")
    task = tasks_db[task_id]
    if task["status"] not in ["running", "paused"]:
        raise HTTPException(status_code=400, detail="Only running or paused tasks can be stopped")
    task["status"] = "stopped"
    task["logs"].append(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] 任务被用户手动停止。")
    if task_id in tasks_events:
        tasks_events[task_id].set()
    return {"status": "stopped"}


@app.get("/api/history")
async def get_history() -> List[Dict[str, Any]]:
    """获取所有历史解析成功的任务列表，包含动态扫描的文件夹结果和正在进行的任务"""
    history_file = os.path.join(DATA_DIR, "task_history.json")
    history_list = []
    if os.path.exists(history_file):
        try:
            with open(history_file, "r", encoding="utf-8") as f:
                history_list = json.load(f)
        except Exception:
            pass

    for item in history_list:
        if "status" not in item:
            item["status"] = "completed"

    scanned_items = []
    if os.path.exists(OUTPUT_DIR):
        for folder_name in os.listdir(OUTPUT_DIR):
            folder_path = os.path.join(OUTPUT_DIR, folder_name)
            if os.path.isdir(folder_path):
                image_paths_file = os.path.join(folder_path, "image_paths.json")
                if os.path.exists(image_paths_file):
                    exists = False
                    for item in history_list:
                        if item.get("task_id") == folder_name or os.path.basename(item.get("output_dir", "")) == folder_name:
                            exists = True
                            break
                    if not exists:
                        is_upload = folder_name.startswith("upload_")
                        ts_str = folder_name.replace("upload_", "")
                        try:
                            if is_upload:
                                ts = float(ts_str)
                                if ts > 1000000000000:
                                    ts /= 1000.0
                                dt_obj = datetime.datetime.fromtimestamp(ts)
                            else:
                                dt_obj = datetime.datetime.strptime(folder_name, "%Y%m%d_%H%M%S")
                            dt = dt_obj.strftime("%Y-%m-%d %H:%M:%S")
                        except Exception:
                            try:
                                mtime = os.path.getmtime(image_paths_file)
                                dt = datetime.datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M:%S")
                            except Exception:
                                dt = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        
                        task_type = "upload" if is_upload else "url"
                        name = "本地文件解析" if is_upload else "网页链接解析"
                        
                        scanned_items.append({
                            "task_id": folder_name,
                            "name": name,
                            "type": task_type,
                            "output_dir": folder_path,
                            "timestamp": dt,
                            "status": "completed"
                        })
    
    if scanned_items:
        history_list.extend(scanned_items)
        try:
            with open(history_file, "w", encoding="utf-8") as f:
                json.dump(history_list, f, ensure_ascii=False, indent=4)
        except Exception as e:
            logger.error(f"Failed to auto-sync task history: {e}")

    active_items = []
    for task_id, task in tasks_db.items():
        status = task.get("status", "pending")
        if status in ["running", "paused", "pending"]:
            if not any(item.get("task_id") == task_id for item in history_list):
                is_upload = task_id.startswith("upload_") or "file_path" in task or "file_paths" in task
                name = task.get("merged_url") or (task.get("file_paths")[0] if task.get("file_paths") else "未知文件")
                if len(name) > 60:
                    name = name[:57] + "..."
                
                task_type = "upload" if is_upload else "url"
                display_name = f"[进行中] {name}" if status == "running" else f"[暂停] {name}"
                
                dt = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                
                active_items.append({
                    "task_id": task_id,
                    "name": display_name,
                    "type": task_type,
                    "output_dir": task.get("output_dir", ""),
                    "timestamp": dt,
                    "status": status
                })
                
    return history_list + active_items


@app.delete("/api/history/{task_id}")
async def delete_history_item(task_id: str) -> Dict[str, Any]:
    """从任务历史记录中删除一条项目"""
    history_file = os.path.join(DATA_DIR, "task_history.json")
    if os.path.exists(history_file):
        try:
            with open(history_file, "r", encoding="utf-8") as f:
                history_list = json.load(f)
            new_list = [item for item in history_list if item.get("task_id") != task_id]
            with open(history_file, "w", encoding="utf-8") as f:
                json.dump(new_list, f, ensure_ascii=False, indent=4)
            return {"status": "success"}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    return {"status": "success"}


@app.post("/api/history/load/{task_id}")
async def load_historic_task(task_id: str, payload: Dict[str, str]) -> Dict[str, Any]:
    """从本地硬盘的输出目录还原历史任务到内存，供前端回放预览和下载"""
    output_dir = payload.get("output_dir")
    if not output_dir or not os.path.exists(output_dir):
        raise HTTPException(status_code=404, detail="Output directory not found")
        
    image_paths_file = os.path.join(output_dir, "image_paths.json")
    if not os.path.exists(image_paths_file):
        raise HTTPException(status_code=400, detail="image_paths.json missing in output directory")
        
    try:
        with open(image_paths_file, "r", encoding="utf-8") as f:
            image_paths = json.load(f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read image_paths.json: {e}")
        
    # Rebuild task dict
    task = {
        "status": "completed",
        "progress": 100,
        "completed_count": len(image_paths),
        "total_count": len(image_paths),
        "logs": ["从历史记录中成功加载。"],
        "error": "",
        "current_processing": set(),
        "image_paths": image_paths,
        "results": {},
        "output_dir": output_dir,
        "lock": threading.Lock()
    }
    
    # Read slices
    for idx, img_path in enumerate(image_paths):
        slice_md = os.path.join(output_dir, f"slice_{idx + 1}.md")
        slice_csv = os.path.join(output_dir, f"table_{idx + 1}.csv")
        if os.path.exists(slice_md):
            try:
                with open(slice_md, "r", encoding="utf-8") as f:
                    content = f.read()
                content = ensure_markdown_table(content)
                rtype = "image" if content == "RETAIN_IMAGE" else ("table" if has_markdown_table(content) else "text")
                res_data = {"type": rtype, "content": content}
                if os.path.exists(slice_csv):
                    res_data["csv_path"] = slice_csv
                task["results"][idx] = (img_path, res_data)
            except Exception:
                pass
                
    # Load final compiled files if they exist on disk
    final_md_file = os.path.join(output_dir, "final_output.md")
    final_txt_file = os.path.join(output_dir, "final_output.txt")
    final_csv_file = os.path.join(output_dir, "final_output.csv")
    final_xlsx_file = os.path.join(output_dir, "final_output.xlsx")
    
    if os.path.exists(final_md_file):
        try:
            with open(final_md_file, "r", encoding="utf-8") as f:
                task["final_md"] = f.read()
        except Exception:
            task["final_md"] = ""
    else:
        task["final_md"] = ""
        
    if os.path.exists(final_txt_file):
        try:
            with open(final_txt_file, "r", encoding="utf-8") as f:
                task["final_txt"] = f.read()
        except Exception:
            task["final_txt"] = ""
    else:
        task["final_txt"] = ""
        
    if os.path.exists(final_csv_file):
        try:
            with open(final_csv_file, "rb") as f:
                task["csv_bytes"] = f.read()
        except Exception:
            pass
            
    if os.path.exists(final_xlsx_file):
        try:
            with open(final_xlsx_file, "rb") as f:
                task["excel_bytes"] = f.read()
        except Exception:
            pass
            
    tasks_db[task_id] = task
    tasks_events[task_id] = asyncio.Event()
    return {"status": "success", "task_id": task_id}


@app.get("/api/download/{task_id}/{file_type}")
def download_task_file(task_id: str, file_type: str) -> Response:
    """下载任务结果文件"""
    if task_id not in tasks_db:
        raise HTTPException(status_code=404, detail="Task not found")

    task = tasks_db[task_id]
    if task["status"] == "running":
        raise HTTPException(status_code=400, detail="Task is still running")

    download_map = {
        "markdown": (task["final_md"].encode("utf-8"), "text/markdown", f"output_{task_id}.md"),
        "txt": (task["final_txt"].encode("utf-8"), "text/plain", f"output_{task_id}.txt"),
        "csv": (task["csv_bytes"], "text/csv", f"output_{task_id}.csv"),
        "excel": (task["excel_bytes"],
                  "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                  f"output_{task_id}.xlsx"),
        "html": (md_to_html(task["final_md"]).encode("utf-8"), "text/html", f"output_{task_id}.html"),
        "latex": (md_to_latex(task["final_md"]).encode("utf-8"), "application/x-latex", f"output_{task_id}.tex"),
        "docx": (md_to_docx(task["final_md"]),
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                 f"output_{task_id}.docx"),
    }

    if file_type in download_map:
        content, media_type, filename = download_map[file_type]
        if content is None:
            raise HTTPException(status_code=404, detail="No data available for this format")
        return Response(
            content=content,
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    if file_type == "zip":
        return _create_zip_download(task_id, task)

    raise HTTPException(status_code=400, detail="Invalid file type")


def _create_zip_download(task_id: str, task: Dict[str, Any]) -> Response:
    """创建 ZIP 打包下载"""
    output_dir = task["output_dir"]
    if not output_dir or not os.path.exists(output_dir):
        raise HTTPException(status_code=404, detail="Output directory not found")

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        report_files = [
            "final_output.md", "final_output.txt", "final_output.html",
            "final_output.tex", "final_output.docx", "final_output.csv",
            "final_output.xlsx"
        ]
        for f_name in report_files:
            f_path = os.path.join(output_dir, f_name)
            if os.path.exists(f_path):
                zf.write(f_path, f_name)

    zip_buffer.seek(0)
    return Response(
        content=zip_buffer.getvalue(),
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=extracted_files_{task_id}.zip"}
    )


@app.put("/api/result/{task_id}/{index}")
async def update_result(task_id: str, index: int, data: Dict[str, Any] = Body(...)) -> Dict[str, Any]:
    """允许前端编辑修改某个切片的识别内容和格式"""
    if task_id not in tasks_db:
        raise HTTPException(status_code=404, detail="Task not found")
    task = tasks_db[task_id]
    if index not in task["results"]:
        raise HTTPException(status_code=404, detail="Slice not found")

    new_content = data.get("content", "")
    new_content = re.sub(r"<br\s*/?>", " ", new_content, flags=re.IGNORECASE)
    forced_type = data.get("type")
    img_path, res_data = task["results"][index]

    if forced_type in ["text", "table", "image", "error", "ignore"]:
        new_type = forced_type
    else:
        if new_content.strip() == "RETAIN_IMAGE":
            new_type = "image"
        elif has_markdown_table(new_content):
            new_type = "table"
        else:
            new_type = "text"

    with task["lock"]:
        task["results"][index] = (img_path, {"type": new_type, "content": new_content})

    output_dir = task["output_dir"]
    if output_dir and os.path.exists(output_dir):
        out_md = os.path.join(output_dir, f"slice_{index + 1}.md")
        with open(out_md, "w", encoding="utf-8") as f:
            f.write(new_content)

    # Re-compile outputs to update ZIP and exports on slice edit/delete
    try:
        compile_task_outputs(task_id)
    except Exception as e:
        logger.error(f"Error recompiling outputs after result update: {e}")

    return {"status": "success"}


@app.post("/api/reextract/{task_id}/{index}")
async def reextract_slice(task_id: str, index: int, data: Dict[str, Any] = Body(None)) -> Dict[str, Any]:
    """重新使用配置对特定切片进行提取识别"""
    if task_id not in tasks_db:
        raise HTTPException(status_code=404, detail="Task not found")
    task = tasks_db[task_id]

    if not task.get("image_paths") or index < 0 or index >= len(task["image_paths"]):
        raise HTTPException(status_code=404, detail="Slice or Image path not found")

    image_path = task["image_paths"][index]
    config = task.get("config", {"api_type": "Local OCR", "api_key": "", "base_url": "", "model_name": ""})
    output_dir = task.get("output_dir")
    if not output_dir or not os.path.exists(output_dir):
        raise HTTPException(status_code=400, detail="Output directory not found")

    formats = data.get("formats") if data else None
    if not formats:
        formats = task.get("formats")

    try:
        analyzer = ImageAnalyzer(
            api_type=config["api_type"],
            api_key=config["api_key"],
            base_url=config["base_url"],
            model_name=config["model_name"],
            split_len=config.get("split_len"),
            split_regex=config.get("split_regex")
        )
        output_csv = os.path.join(output_dir, f"table_{index + 1}.csv")

        loop = asyncio.get_running_loop()
        res = await loop.run_in_executor(
            None,
            lambda: analyzer.analyze(image_path, output_csv_path=output_csv, formats=formats)
        )

        out_md = os.path.join(output_dir, f"slice_{index + 1}.md")
        with open(out_md, "w", encoding="utf-8") as f:
            f.write(res["content"] if res["type"] in ["text", "table"] else "RETAIN_IMAGE")

        with task["lock"]:
            if res.get("type") == "table" and os.path.exists(output_csv):
                res["csv_path"] = output_csv
            task["results"][index] = (image_path, res)

        compile_task_outputs(task_id)

        return {"status": "success", "type": res["type"], "content": res["content"]}
    except Exception as e:
        logger.error(f"Failed to re-extract slice {index + 1}: {e}")
        raise HTTPException(status_code=500, detail=f"重新提取失败: {str(e)}")


@app.post("/api/compile/{task_id}")
async def recompile_task(task_id: str) -> Dict[str, Any]:
    """重新合并所有切片结果"""
    if task_id not in tasks_db:
        raise HTTPException(status_code=404, detail="Task not found")
    compile_task_outputs(task_id)
    task = tasks_db[task_id]
    has_table = task["csv_bytes"] is not None
    return {"status": "success", "has_table": has_table}


@app.put("/api/compiled_result/{task_id}")
async def update_compiled_result(task_id: str, payload: Dict[str, str] = Body(...)) -> Dict[str, Any]:
    """更新最终合并的 Markdown 内容并重新生成本地合并文件"""
    if task_id not in tasks_db:
        raise HTTPException(status_code=404, detail="Task not found")
        
    task = tasks_db[task_id]
    new_md = payload.get("final_md", "")
    task["final_md"] = new_md
    
    # Clean basic MD formats for plain text output
    task["final_txt"] = new_md.replace("#", "").replace("*", "").replace("`", "")
    
    # Save outputs to disk (re-generates docx, html, tex, zip etc)
    output_dir = task.get("output_dir")
    if output_dir and os.path.exists(output_dir):
        _save_outputs_to_disk(task, output_dir)
        
    return {"status": "success"}


# ── 启动入口 ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    
    cyan = "\033[96m"
    yellow = "\033[93m"
    green = "\033[92m"
    magenta = "\033[95m"
    reset = "\033[0m"

    banner = f"""
┌────────────────────────────────────────────────────────────────────────┐
│                                                                        │
│   {cyan}██████  ███    ███ ███    ██ ██  ███████ ██   ██ ████████  ██████{reset}    │
│  {cyan}██    ██ ████  ████ ████   ██ ██  ██       ██ ██     ██    ██    ██{reset}   │
│  {cyan}██    ██ ██ ████ ██ ██ ██  ██ ██  █████     ███      ██    ██████{reset}    │
│  {cyan}██    ██ ██  ██  ██ ██  ██ ██ ██  ██       ██ ██     ██    ██   ██{reset}   │
│   {cyan}██████  ██      ██ ██   ████ ██  ███████ ██   ██    ██    ██    ██{reset}  │
│                                                                        │
│                   {green}OmniExtract 万象多模态提取引擎 v1.4{reset}                  │
│                       {yellow}http://127.0.0.1:8000{reset}                            │
│                                                                        │
│             {magenta}🌐 服务端口:{reset}  http://127.0.0.1:8000                      │
│             {magenta}📂 输出目录:{reset}  data/output                                │
│             {magenta}🚀 服务框架:{reset}  FastAPI 极速引擎                           │
│                                                                        │
└────────────────────────────────────────────────────────────────────────┘
"""
    print(banner)
    uvicorn.run(app, host="0.0.0.0", port=8000)
